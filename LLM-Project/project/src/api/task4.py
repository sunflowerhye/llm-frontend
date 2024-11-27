from flask import Blueprint, request, jsonify, send_file
from common import call_openai_api
from docx import Document
from io import BytesIO

task4 = Blueprint('task4', __name__)

@task4.route('/generate-indoor-event-plan', methods=['POST'])
def generate_indoor_event_plan():
    try:
        data = request.json
        space_and_budget = data.get('spaceAndBudget', '10평, 200만 원')

        # 장소 추천 프롬프트
        location_prompt = (
            f"다음 조건에 맞는 실내 이벤트 장소를 추천해 주세요. 지역 이름(예: 성수, 홍대, 가로수길 등)만 제시하세요:\n"
            f"1. 공간과 예산: {space_and_budget}\n"
            f"2. 주요 고객 관심사: {data.get('customerInterest', '친환경 제품')}\n"
            f"3. 이벤트 주제: {data.get('theme', '자연 친화적인 라이프스타일')}\n"
        )
        location_recommendation = call_openai_api([{"role": "user", "content": location_prompt}])

        # 이벤트 기획 프롬프트
        prompt = (
            f"다음은 기초 화장품 관련 실내 부스 이벤트 기획에 대한 정보입니다:\n"
            f"1. 목적: {data.get('goal', '브랜드 인지도 향상')}\n"
            f"2. 타겟층: {data.get('targetAudience', '20~30대 여성 고객')}\n"
            f"3. 공간과 예산: {space_and_budget}\n"
            f"4. 주요 고객 관심사: {data.get('customerInterest', '친환경 제품')}\n"
            f"5. 이벤트 주제: {data.get('theme', '자연 친화적인 라이프스타일')}\n"
            f"6. 추천 장소: {location_recommendation}\n\n"
            f"위 내용을 바탕으로 체계적이고 창의적인 실내 부스 이벤트 기획서를 작성해 주세요. "
            f"결과는 다음 항목을 포함해 주세요:\n"
            f"1. 이벤트 개요\n"
            f"2. 부스 디자인 및 구성\n"
            f"3. 주요 활동 세부 내용\n"
            f"4. 기대 효과"
        )
        event_plan = call_openai_api([{"role": "user", "content": prompt}])
        return jsonify({'eventPlan': event_plan})

    except Exception as e:
        print(f"Error in /generate-indoor-event-plan: {e}")
        return jsonify({'error': f"서버 오류 발생: {e}"}), 500


@task4.route('/generate-outdoor-event-plan', methods=['POST'])
def generate_outdoor_event_plan():
    try:
        data = request.json
        space_and_budget = data.get('spaceAndBudget', '소형 천막, 100만 원')

        # 장소 추천 프롬프트
        location_prompt = (
            f"다음 조건에 맞는 실외 이벤트 장소를 추천해 주세요. 지역 이름(예: 성수, 홍대, 가로수길 등)만 제시하세요:\n"
            f"1. 공간과 예산: {space_and_budget}\n"
            f"2. 주요 고객 관심사: {data.get('customerInterest', 'SNS 활동')}\n"
            f"3. 이벤트 주제: {data.get('theme', '활기찬 여름 테마')}\n"
        )
        location_recommendation = call_openai_api([{"role": "user", "content": location_prompt}])

        # 이벤트 기획 프롬프트
        prompt = (
            f"다음은 기초 화장품 관련 실외 부스 이벤트 기획에 대한 정보입니다:\n"
            f"1. 목적: \n{data.get('goal', '브랜드 홍보')}\n"
            f"2. 타겟층: \n{data.get('targetAudience', '20~30대 젊은 소비자')}\n"
            f"3. 공간과 예산: \n{space_and_budget}\n"
            f"4. 주요 고객 관심사: \n{data.get('customerInterest', 'SNS 활동')}\n"
            f"5. 이벤트 주제: \n{data.get('theme', '활기찬 여름 테마')}\n"
            f"6. 추천 장소: \n{location_recommendation}\n\n"
            f"위 내용을 바탕으로 체계적이고 창의적인 실외 부스 이벤트 기획서를 작성해 주세요. "
            f"결과는 다음 항목을 포함해 주세요:\n"
            f"1. 이벤트 개요\n"
            f"2. 부스 디자인 및 구성\n"
            f"3. 주요 활동 세부 내용\n"
            f"4. 기대 효과"
        )
        event_plan = call_openai_api([{"role": "user", "content": prompt}])
        return jsonify({'eventPlan': event_plan})

    except Exception as e:
        print(f"Error in /generate-outdoor-event-plan: {e}")
        return jsonify({'error': f"서버 오류 발생: {e}"}), 500

#타임 테이블 생성
@task4.route('/download-timetable', methods=['POST'])
def download_timetable():
    try:
        data = request.json
        space_and_budget = data.get('spaceAndBudget', '10평, 200만 원')
        environment = data.get('environment', '실내')

        # 장소 추천 프롬프트
        location_prompt = (
            f"다음 조건에 맞는 {environment} 이벤트 장소를 추천해 주세요(넓은 지역 이름만 제공):\n"
            f"1. 공간과 예산: {space_and_budget}\n"
            f"2. 주요 고객 관심사: {data.get('customerInterest', '친환경 제품')}\n"
            f"3. 이벤트 주제: {data.get('theme', '자연 친화적인 라이프스타일')}\n"
        )
        location_recommendation = call_openai_api([{"role": "user", "content": location_prompt}])

        # 워드 문서 생성
        doc = Document()
        doc.add_heading('이벤트 기획서', level=1)
        doc.add_paragraph(f"목적: {data.get('goal', '브랜드 홍보')}")
        doc.add_paragraph(f"타겟층: {data.get('targetAudience', '20~30대 소비자')}")
        doc.add_paragraph(f"공간과 예산: {space_and_budget}")
        doc.add_paragraph(f"주요 고객 관심사: {data.get('customerInterest', '친환경 제품')}")
        doc.add_paragraph(f"이벤트 주제: {data.get('theme', '자연 친화적인 라이프스타일')}")
        doc.add_paragraph(f"추천 장소: {location_recommendation}")
        doc.add_paragraph(f"환경: {environment}")
        doc.add_heading('타임 테이블', level=2)

        # 타임 테이블 표 생성
        table = doc.add_table(rows=10, cols=2)  
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = "시간"
        headers[1].text = "활동"

        # 시간대별 일정
        schedule = [
            ("11:00 - 12:00", " "),
            ("12:00 - 13:00", " "),
            ("13:00 - 14:00", " "),
            ("14:00 - 15:00", " "),
            ("15:00 - 16:00", " "),
            ("16:00 - 17:00", " "),
            ("17:00 - 18:00", " "),
            ("18:00 - 19:00", " "),
            ("19:00 - 20:00", " "),
        ]

        for i, (time, activity) in enumerate(schedule, start=1):
            row = table.rows[i].cells
            row[0].text = time
            row[1].text = activity

        # 파일 저장 및 반환
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='event_timetable.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error in /download-timetable: {e}")
        return jsonify({'error': f"서버 오류 발생: {e}"}), 500