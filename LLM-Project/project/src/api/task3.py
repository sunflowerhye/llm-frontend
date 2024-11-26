from flask import Blueprint, request, jsonify, send_file
from common import call_openai_api, data, find_best_match
from io import BytesIO
from docx import Document

task3 = Blueprint('task3', __name__)

#Task3 기획서 작성
@task3.route('/generate-marketing-plan', methods=['POST'])
def generate_marketing_plan():
    try:
        input_data = request.json
        goal = input_data.get('goal', '')
        strategy = input_data.get('strategy', '')
        target_audience = input_data.get('targetAudience', '')
        budget = input_data.get('budget', '')

        # 데이터셋에서 제품명과 목표 비교
        product_match, product_score = find_best_match(goal, '제품명')  # 가장 유사한 하나의 제품 선택

        # 프롬프트 생성
        prompt = (
            f"당신은 기초 화장품 브랜드 홍보 기획안 작성 전문가입니다. "
            f"아래 정보를 바탕으로 체계적이고 구체적인 기초 화장품 브랜드 홍보 기획안을 작성해주세요. "
            f"각 항목은 번호로 구분하여 작성해주세요. "
            f"모든 항목은 반드시 존댓말로 작성해야 하며, 응답 내용은 공백 제외 700자 이상이어야 합니다. "
            f"아래 항목들을 포함하되, 기획안의 제목은 작성하지 말아주세요:\n\n"
            f"1. 목표: \n{goal}\n\n"
            f"2. 타겟층: \n{target_audience}\n\n"
            f"3. 전략: \n{strategy}\n\n"
            f"4. 홍보: \n제품 홍보를 위한 세부 계획을 작성해주세요 (예: 소셜 미디어, 광고, 협찬 등).\n\n"
            f"5. 예산: \n{budget}에 기반한 예산 항목을 나열해주세요 (예: 장소 대여, 장비 구매 등).\n\n"
            f"6. 실행: \n제품 홍보 행사와 관련된 구체적인 실행 계획을 작성해주세요 (예: 진행 방식 등).\n\n"
            f"결론: 마케팅과 홍보 활동의 중요성을 강조해서 작성해주세요. "
            f"이 계획의 성공을 위한 주요 요인들을 포함하여 설득력 있는 결론을 작성해주세요."
        )

        messages = [{"role": "user", "content": prompt}]
        marketing_plan = call_openai_api(messages, max_tokens=1200)  # 토큰 수 확장

        # 추천 상품 추가
        if product_match:
            recommendations_text = f"\n\n추천 제품:\n- {product_match}"
            marketing_plan += recommendations_text

        # 결과 반환
        return jsonify({'marketingPlan': marketing_plan})

    except Exception as e:
        print(f"Error in /generate-marketing-plan endpoint: {e}")
        return jsonify({'error': f"서버 오류 발생: {e}"}), 500

#Task3 기획서 워드로 다운로드
@task3.route('/download-marketing-plan', methods=['POST'])
def download_marketing_plan():
    try:
        input_data = request.json
        goal = input_data.get('goal', '')
        strategy = input_data.get('strategy', '')
        target_audience = input_data.get('targetAudience', '')
        budget = input_data.get('budget', '')
        execution = input_data.get('execution', '')
        conclusion = input_data.get('conclusion', '')

        # Create a Word document
        doc = Document()
        doc.add_heading('이벤트 기획서', level=1)

        # Add metadata
        doc.add_paragraph("작성일자: 2024년 11월 23일")
        doc.add_paragraph("작성자: 마케팅 팀")

        # Add table for '이벤트 개요'
        doc.add_heading('1. 이벤트 개요', level=2)
        table1 = doc.add_table(rows=3, cols=2)
        table1.style = 'Table Grid'
        table1.cell(0, 0).text = "목적"
        table1.cell(0, 1).text = goal or "홍보 및 매출 증대"
        table1.cell(1, 0).text = "이벤트 방식"
        table1.cell(1, 1).text = strategy or "소셜 미디어 캠페인"
        table1.cell(2, 0).text = "기대효과(목표치)"
        table1.cell(2, 1).text = "브랜드 인지도 상승 및 신규 고객 유치"

        # Add table for '이벤트 진행계획'
        doc.add_heading('2. 이벤트 진행계획', level=2)
        table2 = doc.add_table(rows=8, cols=2)
        table2.style = 'Table Grid'
        rows_data = [
            ("제목", "신제품 홍보 이벤트"),
            ("참여대상", target_audience or "20-30대 여성"),
            ("이벤트방법", strategy or "소셜 미디어 광고 및 협찬"),
            ("참여인원(예상)", "100명"),
            ("이벤트기간", "4주"),
            ("경품종류(혜택)", "제품 할인 쿠폰"),
            ("경품지급방법", "이메일 전송"),
            ("예산비용", budget or "100만 원")
        ]
        for i, (label, value) in enumerate(rows_data):
            table2.cell(i, 0).text = label
            table2.cell(i, 1).text = value

        # Add table for '마케팅 방안'
        doc.add_heading('3. 마케팅 방안', level=2)
        table3 = doc.add_table(rows=2, cols=2)
        table3.style = 'Table Grid'
        table3.cell(0, 0).text = "내부 홍보"
        table3.cell(0, 1).text = "사내 이메일 공지"
        table3.cell(1, 0).text = "외부 홍보"
        table3.cell(1, 1).text = strategy or "SNS 캠페인, 유튜브 광고"

       
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        
        return send_file(file_stream, as_attachment=True, download_name="event_plan.docx")

    except Exception as e:
        print(f"Error in /download-marketing-plan endpoint: {e}")
        return jsonify({'error': f"서버 오류 발생: {e}"}), 500
